import streamlit as st
import os
import re
import shutil
import tempfile
from zipfile import ZipFile, ZIP_DEFLATED
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import subprocess

# Configurazione pagina
st.set_page_config(
    page_title="Generatore Buoni Carburante",
    page_icon="⛽",
    layout="centered"
)

# CSS personalizzato
st.markdown("""
<style>
    .main-title {
        text-align: center;
        color: #1E40AF;
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    .subtitle {
        text-align: center;
        color: #6B7280;
        font-size: 1.1rem;
        margin-bottom: 2rem;
    }
    .success-box {
        background-color: #ECFDF5;
        border: 2px solid #10B981;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Titolo
st.markdown('<h1 class="main-title">🎫 Generatore Buoni Carburante</h1>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">Genera i tuoi buoni in pochi click</p>', unsafe_allow_html=True)

# Session state
if 'step' not in st.session_state:
    st.session_state.step = 1

# Funzioni
def convert_svg_to_png(svg_data, output_path):
    try:
        import cairosvg
        cairosvg.svg2png(bytestring=svg_data, write_to=output_path)
        return True
    except Exception as e:
        st.error(f"Errore conversione SVG: {e}")
        return False

def extract_voucher_number(folder_name):
    match = re.search(r'- ([a-zA-Z0-9]+) -', folder_name)
    if match:
        return match.group(1)
    match = re.search(r'-([a-zA-Z0-9]+)-', folder_name)
    return match.group(1) if match else None

def process_buoni(template_file, zip_file, progress_bar):
    temp_dir = tempfile.mkdtemp()
    qr_folder = os.path.join(temp_dir, 'qr_convertiti')
    output_folder = os.path.join(temp_dir, 'buoni_generati')
    os.makedirs(qr_folder)
    os.makedirs(output_folder)
    
    try:
        template_path = os.path.join(temp_dir, 'template.docx')
        with open(template_path, 'wb') as f:
            f.write(template_file.getbuffer())
        
        zip_path = os.path.join(temp_dir, 'qrcodes.zip')
        with open(zip_path, 'wb') as f:
            f.write(zip_file.getbuffer())
        
        vouchers = []
        with ZipFile(zip_path, 'r') as zf:
            svg_files = [f for f in zf.namelist() if f.lower().endswith('.svg')]
            
            for svg_path in svg_files:
                folder = svg_path.split('/')[0]
                num = extract_voucher_number(folder)
                
                if num and not any(v['number'] == num for v in vouchers):
                    svg_data = zf.read(svg_path)
                    png_path = os.path.join(qr_folder, f'qr_{num}.png')
                    
                    if convert_svg_to_png(svg_data, png_path):
                        vouchers.append({'number': num, 'qr_png': png_path})
        
        total = len(vouchers)
        
        for idx, voucher in enumerate(vouchers):
            progress_bar.progress((idx + 1) / total, 
                                 f"Generazione {idx + 1}/{total}: {voucher['number']}")
            
            doc = Document(template_path)
            
            # Sostituisci numero
            for para in doc.paragraphs:
                if 'Buono n.' in para.text and 'valido' in para.text:
                    parti = para.text.split('Buono n.')
                    if len(parti) >= 2:
                        dopo = parti[1]
                        if 'valido' in dopo:
                            nuovo = f"Buono n. {voucher['number']} valido" + dopo.split('valido', 1)[1]
                            for run in para.runs:
                                run.text = ''
                            if para.runs:
                                para.runs[0].text = nuovo
                            break
            
            # Sostituisci QR
            if len(doc.paragraphs) > 5:
                para = doc.paragraphs[5]
                for run in list(para.runs):
                    drawings = run._element.findall('.//' + qn('w:drawing'))
                    if drawings:
                        run._element.getparent().remove(run._element)
                
                new_run = para.add_run()
                new_run.add_picture(voucher['qr_png'], width=Inches(1.5))
            
            output_path = os.path.join(output_folder, f'Buono_{voucher["number"]}.docx')
            doc.save(output_path)
        
        return output_folder, len(vouchers)
        
    except Exception as e:
        st.error(f"Errore: {e}")
        return None, 0

def create_pdf(output_folder, progress_bar):
    try:
        pdf_files = []
        docx_files = sorted([f for f in os.listdir(output_folder) if f.endswith('.docx')])
        
        for idx, docx_file in enumerate(docx_files):
            progress_bar.progress((idx + 1) / len(docx_files),
                                 f"Conversione PDF {idx + 1}/{len(docx_files)}")
            
            docx_path = os.path.join(output_folder, docx_file)
            
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', output_folder, docx_path
            ], capture_output=True, timeout=30)
            
            pdf_path = docx_path.replace('.docx', '.pdf')
            if os.path.exists(pdf_path):
                pdf_files.append(pdf_path)
        
        from PyPDF2 import PdfMerger
        merger = PdfMerger()
        for pdf in sorted(pdf_files):
            merger.append(pdf)
        
        output_pdf = os.path.join(output_folder, 'buoni_COMPLETI.pdf')
        merger.write(output_pdf)
        merger.close()
        
        return output_pdf
        
    except Exception as e:
        st.error(f"Errore PDF: {e}")
        return None

def create_zip(output_folder):
    zip_path = os.path.join(output_folder, 'buoni_generati.zip')
    
    with ZipFile(zip_path, 'w', ZIP_DEFLATED) as zipf:
        for file in os.listdir(output_folder):
            if file.endswith('.docx'):
                zipf.write(os.path.join(output_folder, file), file)
    
    return zip_path

# INTERFACCIA
if st.session_state.step == 1:
    st.markdown("### 📤 Carica i File")
    
    col1, col2 = st.columns(2)
    
    with col1:
        template_file = st.file_uploader("📄 Template Word", type=['docx'])
    
    with col2:
        zip_file = st.file_uploader("📦 ZIP QR Codes", type=['zip'])
    
    if template_file and zip_file:
        st.success("✅ File caricati!")
        
        if st.button("🚀 Genera Buoni", type="primary", use_container_width=True):
            st.session_state.template = template_file
            st.session_state.zip = zip_file
            st.session_state.step = 2
            st.rerun()

elif st.session_state.step == 2:
    st.markdown("### ⚙️ Elaborazione...")
    
    progress_bar = st.progress(0, "Inizio...")
    
    output_folder, count = process_buoni(
        st.session_state.template,
        st.session_state.zip,
        progress_bar
    )
    
    if output_folder and count > 0:
        st.session_state.output_folder = output_folder
        st.session_state.buoni_count = count
        st.session_state.step = 3
        st.rerun()
    else:
        st.error("❌ Errore generazione")
        if st.button("Riprova"):
            st.session_state.step = 1
            st.rerun()

elif st.session_state.step == 3:
    st.markdown('<div class="success-box">', unsafe_allow_html=True)
    st.markdown(f"### ✅ {st.session_state.buoni_count} Buoni Generati!")
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown("### 📥 Scegli Formato")
    
    formato = st.radio(
        "Formato:",
        ["PDF unico (per stampa)", "Word separati (ZIP)"],
        label_visibility="collapsed"
    )
    
    if st.button("⬇️ Scarica", type="primary", use_container_width=True):
        progress = st.progress(0, "Preparazione...")
        
        if "PDF" in formato:
            pdf_path = create_pdf(st.session_state.output_folder, progress)
            
            if pdf_path and os.path.exists(pdf_path):
                with open(pdf_path, 'rb') as f:
                    st.download_button(
                        "📄 Scarica PDF",
                        f.read(),
                        file_name="buoni_carburante.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
        else:
            zip_path = create_zip(st.session_state.output_folder)
            
            if os.path.exists(zip_path):
                with open(zip_path, 'rb') as f:
                    st.download_button(
                        "📦 Scarica ZIP",
                        f.read(),
                        file_name="buoni_carburante.zip",
                        mime="application/zip",
                        use_container_width=True
                    )
    
    if st.button("🔄 Genera Altri"):
        st.session_state.step = 1
        st.rerun()

st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: #9CA3AF;'>⛽ Generatore Buoni Carburante</p>",
    unsafe_allow_html=True
)
